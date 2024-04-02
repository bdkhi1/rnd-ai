from docx import Document


def update_doc(docx_file, text_file):
    # Open the Word document
    doc = Document(docx_file)

    # Read text from the text file
    with open(text_file, 'r') as file:
        text = file.read()

    initial_content = list(doc.paragraphs)

    # Update the content in the Word document
    for paragraph in initial_content:
        if "Outdated content" in paragraph.text:
            paragraph.text = paragraph.text.replace("Outdated content", text)

    # Save the updated Word document
    doc.save(docx_file)
    print("Document updated.\n")


def generate_doc(text_file):
    doc = Document()

    # Read text from the text file
    with open(text_file, 'r') as file:
        text = file.read()

    # Add text to the doc
    doc.add_paragraph(text)

    # Save the doc
    doc.save('generated_doc.docx')
    print("Generated docx file named generated_doc.")


def print_text_file(text_file):
    with open(text_file, 'r') as file:
        text = file.read()

    print(text + "\n")


# Print contents of text file
print("Printing contents of the text file:\n")
print_text_file('text.txt')

# Update the Word doc
print("Updating word document")
update_doc('update_doc.docx', 'text.txt')

# Generate a Word doc
print("Generating word document")
generate_doc('text.txt')
